from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "NUERS_Live_API_Transaction_Test_Report.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "DADCE0"
INK = "111827"
MUTED = "4B5563"
GREEN = "166534"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, color=INK, size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = "Heading {}".format(level)
    run = paragraph.add_run(text)
    run.bold = True
    return paragraph


def add_body(doc, text, after=6):
    paragraph = doc.add_paragraph()
    paragraph.style = "Normal"
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.10
    paragraph.add_run(text)
    return paragraph


def add_callout(doc, title, body, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(4)
    title_run = paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(10)
    title_run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    paragraph.add_run("\n")
    body_run = paragraph.add_run(body)
    body_run.font.name = "Calibri"
    body_run.font.size = Pt(9)
    body_run.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_widths(table, [1.7, 4.8])
    for idx, (label, value) in enumerate(rows):
        left = table.cell(idx, 0)
        right = table.cell(idx, 1)
        set_cell_shading(left, LIGHT_BLUE)
        set_cell_text(left, label, bold=True, color=DARK_BLUE)
        set_cell_text(right, value, size=9)
    return table


def add_matrix_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_widths(table, widths)
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_text(cell, header, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_borders(cells[idx])
            set_cell_text(cells[idx], value, size=9, color=GREEN if str(value).lower() in {"passed", "revoked", "inactive"} else INK)
            if idx == 0:
                set_cell_shading(cells[idx], LIGHT_GRAY)
    return table


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(doc):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("NUERS Live API Transaction Test Report")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def build():
    doc = Document()
    configure_styles(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("NUERS Live API Transaction Test Report")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    sub_run = subtitle.add_run("Successful production API insertion and verification for the Business Account integration endpoint")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor.from_string(MUTED)

    add_key_value_table(doc, [
        ("Live site", "https://nuers.net"),
        ("API endpoint tested", "POST https://nuers.net/public/api/integrations/transactions"),
        ("Business Account", "Servxbit I.T Solutions"),
        ("Test date", "June 3, 2026, Asia/Manila"),
        ("Prepared by", "Codex API test run"),
        ("Security note", "No API secret, bearer token, or password is included in this report."),
    ])

    add_callout(
        doc,
        "Executive result",
        "The live NUERS transaction API accepted a marked production test transaction with HTTP 201. The inserted transaction and generated receipt were both verified through the live merchant APIs. The temporary API key created for the test was revoked and confirmed inactive.",
        fill="EEF7EE",
    )

    add_heading(doc, "1. Test Scope", 1)
    add_body(doc, "The test validated that a Business Account API key with the transactions:write scope can insert a live transaction through the NUERS integration endpoint, and that the resulting transaction and receipt appear through the authenticated merchant APIs.")
    add_body(doc, "This was a production test. The record is intentionally marked so it can be found, excluded from analysis, or removed later if required.")

    add_heading(doc, "2. API Contract Used", 1)
    add_key_value_table(doc, [
        ("Method", "POST"),
        ("URL", "https://nuers.net/public/api/integrations/transactions"),
        ("Authentication", "X-NUERS-API-Key header"),
        ("Required fields used", "external_reference, gross_amount, tax_type"),
        ("Tax setup", "vatable, 12 percent VAT rate, VAT inclusive"),
        ("Temporary key handling", "Created for the test, used once, then revoked."),
    ])

    add_heading(doc, "3. Submitted Payload", 1)
    add_body(doc, "The following payload was submitted. The API key itself is intentionally omitted.")
    add_code_block(doc, """{
  "external_reference": "CODEX-LIVE-API-TEST-B21DF9A8",
  "source_system": "codex_live_api_test",
  "channel": "api_test",
  "transaction_type": "sale",
  "payment_method": "cash",
  "document_type": "sales_invoice",
  "gross_amount": 123.45,
  "tax_type": "vatable",
  "vat_rate": 12,
  "vat_inclusive": true,
  "customer_name": "Codex Live API Test Buyer",
  "customer_tin": "000-000-000-000",
  "branch_code": "MAIN",
  "branch_name": "Main Office",
  "branch_type": "main",
  "items": [
    {
      "description": "Codex live API test item",
      "quantity": 1,
      "unit_price": 123.45
    }
  ]
}""")

    add_heading(doc, "4. API Response", 1)
    add_matrix_table(doc, ["Field", "Observed value", "Result"], [
        ("HTTP status", "201 Created", "Passed"),
        ("Duplicate flag", "No duplicate response was returned", "Passed"),
        ("Transaction ID", "429bd850-0318-437c-9fde-363d63ac6c66", "Passed"),
        ("Receipt number", "VI-20260603-WQEBLQC1", "Passed"),
    ], [1.5, 3.7, 1.3])

    add_heading(doc, "5. Verification Evidence", 1)
    add_matrix_table(doc, ["Verification check", "Live endpoint / signal", "Observed result"], [
        ("Merchant transaction lookup", "/public/api/merchant/transactions?search=CODEX-LIVE-API-TEST-B21DF9A8", "1 matching row found"),
        ("Merchant receipt lookup", "/public/api/merchant/receipts?search=VI-20260603-WQEBLQC1", "1 matching row found"),
        ("Temporary API key cleanup", "/public/api/merchant/api-keys", "Status revoked, active no"),
    ], [1.9, 3.3, 1.3])

    add_heading(doc, "6. Security and Cleanup", 1)
    add_body(doc, "The full API secret was available only immediately after key creation and was not written into this report. Temporary local response files containing tokens or secrets were deleted after the test.")
    add_body(doc, "The temporary live API key was revoked after the transaction was verified. Read-only verification confirmed that its status is revoked and active is no.")

    add_heading(doc, "7. Operational Notes", 1)
    add_matrix_table(doc, ["Item", "Value"], [
        ("Record classification", "Production test transaction"),
        ("External reference", "CODEX-LIVE-API-TEST-B21DF9A8"),
        ("Gross amount", "PHP 123.45"),
        ("Recommended handling", "Keep for audit trail, or filter by external_reference/source_system in reports if test records should be excluded."),
        ("Code changes required", "None for transaction insertion. The live API accepted and persisted the transaction successfully."),
    ], [1.9, 4.6])

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build()
    print(DOCX_PATH)
